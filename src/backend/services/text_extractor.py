"""
Text Extractor Service
Extracts plain text from uploaded meeting documents:
- PDF → pypdf
- DOCX → python-docx
- XLSX → openpyxl
- TXT → direct read
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


def extract_text(content: bytes, filename: str, content_type: str = "") -> str:
    """
    Extract plain text from a file's binary content.

    Args:
        content: Raw file bytes
        filename: Original filename (used to detect type if content_type is missing)
        content_type: MIME type (optional fallback)

    Returns:
        Extracted plain text string, empty string on failure.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "pdf" or "pdf" in content_type:
            return _extract_pdf(content)
        elif ext == "docx" or "wordprocessingml" in content_type:
            return _extract_docx(content)
        elif ext == "xlsx" or "spreadsheetml" in content_type:
            return _extract_xlsx(content)
        elif ext in ("txt", "md", "csv"):
            return content.decode("utf-8", errors="replace")
        else:
            # Best-effort: try UTF-8 decode for unknown types
            decoded = content.decode("utf-8", errors="replace")
            # Only return if it looks like readable text (not binary garbage)
            printable_ratio = sum(1 for c in decoded if c.isprintable()) / max(len(decoded), 1)
            if printable_ratio > 0.7:
                return decoded
            return ""
    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", filename, exc)
        return ""


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    return "\n".join(parts)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)
    # Also extract text from any tables in the document
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text and row_text not in parts:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)
