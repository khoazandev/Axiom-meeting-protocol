import os
import subprocess
import sys

def install_package(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    install_package("python-docx")
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
title = doc.add_heading('HỒ SƠ Ý TƯỞNG DỰ ÁN (PITCH DECK)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Dự án: Smart Meeting AI - Nền tảng hội nghị định hướng quy trình & Bảo mật cho doanh nghiệp')
doc.add_paragraph('Cuộc thi: Olympic Phần mềm Nguồn mở (PMNM) 2026 - Chủ đề DX-OS')
doc.add_paragraph('Nhóm thực hiện: [Tên Nhóm]')
doc.add_page_break()

# 1. Đặt Vấn Đề
doc.add_heading('1. Đặt Vấn Đề (Pain points)', level=1)
doc.add_paragraph('Hiện nay các doanh nghiệp (SME) đang gặp 3 vấn đề lớn trong việc hội họp:', style='List Bullet')
doc.add_paragraph('Lãng phí thời gian: Các cuộc họp diễn ra tràn lan, không có Agenda rõ ràng dẫn đến kém hiệu quả.', style='List Bullet')
doc.add_paragraph('Quản lý rời rạc: Sử dụng Zalo, Google Meet, Zoom không đồng nhất. Các Action Items (nhiệm vụ) sau cuộc họp thường bị lãng quên.', style='List Bullet')
doc.add_paragraph('Bảo mật dữ liệu: Dùng các tool bên thứ ba hoặc AI (như ChatGPT) để tóm tắt biên bản họp tiềm ẩn nguy cơ lộ lọt bí mật kinh doanh.', style='List Bullet')

# 2. Giải pháp 
doc.add_heading('2. Giải Pháp - Smart Meeting AI', level=1)
doc.add_paragraph('Xây dựng một nền tảng quản lý hội họp nội bộ On-Premise (tự lưu trữ), giải quyết bài toán hội họp dựa trên khung kiến trúc H-P-D-I của hệ điều hành DX-OS:')

# HPDI
doc.add_heading('2.1. H (Human) - Không gian tương tác', level=2)
doc.add_paragraph('Môi trường số hóa giúp nhân viên gặp gỡ trực tuyến thông qua nền tảng Video Call chất lượng cao, chia sẻ màn hình, bảng trắng ngay trên trình duyệt mà không cần cài đặt ứng dụng ngoài.', style='List Bullet')

doc.add_heading('2.2. P (Process) - Rào chắn quy trình', level=2)
doc.add_paragraph('Quy trình "Trực thuộc": Yêu cầu người tạo cuộc họp bắt buộc khai báo Agenda, mục tiêu. Nếu không có -> không cho lên lịch.', style='List Bullet')
doc.add_paragraph('Quy trình "Sau họp": Tự động xuất biên bản (Meeting Minutes) và chuyển các nhiệm vụ (Action Items) vào hệ thống quản lý Task nội bộ (Kanban).', style='List Bullet')

doc.add_heading('2.3. D (Data) - Sự thật dữ liệu', level=2)
doc.add_paragraph('Lưu trữ toàn bộ file ghi âm, video, tài liệu đính kèm nội bộ.', style='List Bullet')
doc.add_paragraph('Dashboard thống kê cho CEO: Báo cáo số giờ họp mỗi tuần của từng phòng ban, tỷ lệ hoàn thành công việc sau cuộc họp để tối ưu chi phí vận hành.', style='List Bullet')

doc.add_heading('2.4. I (Intelligence) - Trí tuệ AI', level=2)
doc.add_paragraph('AI Speech-to-Text: Tự động nhận diện và bóc băng giọng nói (tiếng Việt) thành văn bản.', style='List Bullet')
doc.add_paragraph('AI Summary: Tự động đọc biên bản và sinh ra tóm tắt ngắn gọn: Ai phát biểu gì, quyết định cuối cùng là gì, ai làm gì tiếp theo.', style='List Bullet')

# 3. Tính năng cốt lõi MVP
doc.add_heading('3. Tính năng cốt lõi (MVP mang đi thi)', level=1)
doc.add_paragraph('Module 1 - Đặt lịch & Phê duyệt: Lên lịch, mời thành viên, nhập Agenda.', style='List Number')
doc.add_paragraph('Module 2 - Không gian họp ảo: Phòng họp Video Call thời gian thực.', style='List Number')
doc.add_paragraph('Module 3 - AI Trợ lý: Xử lý tóm tắt văn bản và bóc băng ghi âm.', style='List Number')
doc.add_paragraph('Module 4 - Dashboard & Báo cáo: Báo cáo thời lượng họp cho quản trị viên.', style='List Number')

# 4. Kiến trúc công nghệ nguồn mở
doc.add_heading('4. Kiến trúc & Công nghệ nguồn mở (Open Source Stack)', level=1)
doc.add_paragraph('Để đáp ứng tiêu chí cao nhất của giải thưởng, dự án cam kết 100% sử dụng mã nguồn mở:')
doc.add_paragraph('Lõi Video/Voice: LiveKit Server WebRTC.', style='List Bullet')
doc.add_paragraph('AI Nhận diện giọng nói (STT): OpenAI Whisper (Bản Open Source tự host) hoặc PhoWhisper.', style='List Bullet')
doc.add_paragraph('AI Tóm tắt văn bản (LLM): Ollama kết hợp với model Llama 3 hoặc Qwen-2.', style='List Bullet')
doc.add_paragraph('Backend & Logic: Python (FastAPI) để xử lý mượt mà luồng AI, kết hợp Node.js (tùy chọn).', style='List Bullet')
doc.add_paragraph('Frontend: Next.js (React) kết hợp TailwindCSS & Shadcn UI cho giao diện chuyên nghiệp.', style='List Bullet')
doc.add_paragraph('Database: PostgreSQL.', style='List Bullet')

# 5. Phân công nhiệm vụ
doc.add_heading('5. Đề xuất Phân công nhiệm vụ (Nhóm 3 người)', level=1)

p1 = doc.add_paragraph()
p1.add_run('Thành viên 1 (Frontend & UI/UX): ').bold = True
p1.add_run('Thiết kế giao diện, code Frontend (Next.js/React). Tích hợp giao diện phòng họp (LiveKit SDK). Xây dựng Dashboard.')

p2 = doc.add_paragraph()
p2.add_run('Thành viên 2 (Backend & AI Integration): ').bold = True
p2.add_run('Viết API (FastAPI/Node.js). Cài đặt và tích hợp AI (Whisper, Ollama). Xử lý luồng tải file âm thanh và trả về văn bản.')

p3 = doc.add_paragraph()
p3.add_run('Thành viên 3 (Database, DevOps & Thuyết trình): ').bold = True
p3.add_run('Thiết kế CSDL (PostgreSQL). Đóng gói Docker. Phụ trách luồng Process (Quy trình tạo lịch họp) và chuẩn bị tài liệu, slide thuyết trình (Pitching).')

# 6. Lộ trình thực hiện
doc.add_heading('6. Lộ trình thực hiện (Roadmap)', level=1)
doc.add_paragraph('Giai đoạn 1 (Tuần 1-2): Dựng khung dự án, thiết kế Database, setup môi trường Frontend và Backend.', style='List Bullet')
doc.add_paragraph('Giai đoạn 2 (Tuần 3-4): Tích hợp phòng họp Video (LiveKit) và làm tính năng Đặt lịch họp theo quy trình.', style='List Bullet')
doc.add_paragraph('Giai đoạn 3 (Tuần 5-6): Xử lý phần "Xương" nhất: Tích hợp AI bóc băng và tóm tắt (Whisper + Ollama).', style='List Bullet')
doc.add_paragraph('Giai đoạn 4 (Tuần 7-8): Hoàn thiện Dashboard, fix bug, tối ưu UI/UX, đóng gói Docker và làm Slide.', style='List Bullet')

doc.save('Smart_Meeting_AI_Y_Tuong.docx')
print("Đã tạo file Smart_Meeting_AI_Y_Tuong.docx thành công!")
